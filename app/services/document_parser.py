from __future__ import annotations

from abc import ABC, abstractmethod
from pathlib import Path

from app.services.pdf_parser import PdfParseError, RoutedPdfParser


class DocumentParseError(PdfParseError):
    """Raised when a supported review document cannot be parsed."""


class DocumentExtractionProvider(ABC):
    """Unified provider interface for review document parsing."""

    provider_name = "unknown"
    provider_label = "Unknown Document Parser"
    is_placeholder = False
    provider_note = ""

    @abstractmethod
    def extract(self, file_path: Path) -> tuple[str, int]:
        """Extract plain text and a logical page/section count."""


class DocxParser(DocumentExtractionProvider):
    """Extract text from a `.docx` file."""

    provider_name = "docx_parser"
    provider_label = "代码版 DOCX 解析"

    def extract(self, file_path: Path) -> tuple[str, int]:
        if not file_path.exists():
            raise DocumentParseError(f"未找到文档文件：{file_path.name}")
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover
            raise DocumentParseError(
                "未安装 python-docx，无法解析 Word 文档。请先执行 `pip install -r requirements.txt`。"
            ) from exc

        try:
            document = Document(str(file_path))
        except Exception as exc:  # pragma: no cover
            raise DocumentParseError(f"{file_path.name} 解析失败：{exc}") from exc

        blocks: list[str] = []
        paragraph_count = 0
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                blocks.append(text)
                paragraph_count += 1

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))
                    paragraph_count += 1

        if not blocks:
            raise DocumentParseError(f"{file_path.name} 未提取到可用文本。")

        return "\n".join(blocks), max(1, paragraph_count)


class RoutedDocumentParser(DocumentExtractionProvider):
    """Route parsing by file suffix and keep the existing PDF pipeline intact."""

    provider_name = "routed_document_parser"
    provider_label = "统一文档解析路由器"

    def __init__(
        self,
        *,
        pdf_parser: RoutedPdfParser | None = None,
        docx_parser: DocxParser | None = None,
    ) -> None:
        self.pdf_parser = pdf_parser or RoutedPdfParser()
        self.docx_parser = docx_parser or DocxParser()
        self.provider_note = ""
        self.last_note = ""
        self.last_snapshot = None

    def extract(self, file_path: Path) -> tuple[str, int]:
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            text, page_count = self.pdf_parser.extract(file_path)
            self.provider_note = getattr(self.pdf_parser, "provider_note", "")
            self.last_note = str(getattr(self.pdf_parser, "last_note", self.provider_note) or "")
            self.last_snapshot = getattr(self.pdf_parser, "last_snapshot", None)
            return text, page_count
        if suffix == ".docx":
            text, paragraph_count = self.docx_parser.extract(file_path)
            self.provider_note = "DOCX 文档已走本地代码解析链路。"
            self.last_note = self.provider_note
            self.last_snapshot = None
            return text, paragraph_count
        if suffix == ".doc":
            raise DocumentParseError("当前暂不直接支持 `.doc`，请先转换为 `.docx` 后再上传。")
        raise DocumentParseError(f"暂不支持的文件类型：{file_path.suffix or '无后缀'}")
